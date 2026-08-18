from __future__ import annotations

from io import BytesIO
import re
import unicodedata
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn


def excel_bytes(df, sheet="结果"):
    b = BytesIO()
    with pd.ExcelWriter(b, engine="openpyxl") as w:
        df.to_excel(w, sheet_name=str(sheet)[:31], index=False)
    return b.getvalue()


def _sanitize_text(text: str) -> str:
    """清理可能导致 Word 缺字/方框的控制字符，同时保留正常中英文、希腊字母和公式符号。"""
    s = unicodedata.normalize("NFC", str(text or ""))
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = s.replace("\u2028", "\n").replace("\u2029", "\n")
    s = s.replace("\ufeff", "")
    # 去除零宽字符与 Word 不接受的 C0 控制字符（保留换行/制表）
    s = re.sub(r"[\u200b\u200c\u200d\u2060]", "", s)
    s = "".join(ch for ch in s if ch in "\n\t" or ord(ch) >= 32)
    return s


def _set_run_font(run, east_asia="微软雅黑", latin="Times New Roman", size=None, bold=None):
    run.font.name = latin
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), latin)


def _apply_paragraph_font(paragraph, size=11, bold=None):
    for run in paragraph.runs:
        _set_run_font(run, size=size, bold=bold)


def _configure_document(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.3)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.3)

    styles = doc.styles
    for style_name, size, bold in [
        ("Normal", 11, False),
        ("Title", 20, True),
        ("Heading 1", 16, True),
        ("Heading 2", 14, True),
        ("Heading 3", 12, True),
        ("List Bullet", 11, False),
        ("List Number", 11, False),
    ]:
        try:
            st = styles[style_name]
            st.font.name = "Times New Roman"
            st.font.size = Pt(size)
            st.font.bold = bold
            st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        except Exception:
            pass


def _add_markdown_paragraph(doc: Document, raw_line: str):
    s = _sanitize_text(raw_line).strip()
    if not s:
        return
    if re.fullmatch(r"[-—_=]{3,}", s):
        return

    if s.startswith("### "):
        p = doc.add_heading(s[4:].strip(), level=3)
        _apply_paragraph_font(p, 12, True)
        return
    if s.startswith("## "):
        p = doc.add_heading(s[3:].strip(), level=2)
        _apply_paragraph_font(p, 14, True)
        return
    if s.startswith("# "):
        p = doc.add_heading(s[2:].strip(), level=1)
        _apply_paragraph_font(p, 16, True)
        return

    if re.match(r"^[-*•]\s+", s):
        content = re.sub(r"^[-*•]\s+", "", s)
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(content.replace("**", ""))
        _set_run_font(r, size=11)
    elif re.match(r"^\d+[\.、]\s*", s):
        content = re.sub(r"^\d+[\.、]\s*", "", s)
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(content.replace("**", ""))
        _set_run_font(r, size=11)
    else:
        p = doc.add_paragraph()
        r = p.add_run(s.replace("**", ""))
        _set_run_font(r, size=11)

    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(5)


def docx_bytes(title, text, sources=None):
    """生成中文兼容的 Word 报告。

    重点修复：显式设置 East Asia 字体，避免部分 Windows Word 环境把中文显示成方框。
    """
    doc = Document()
    _configure_document(doc)

    title_p = doc.add_heading(_sanitize_text(title), 0)
    _apply_paragraph_font(title_p, 20, True)

    for line in _sanitize_text(text).splitlines():
        _add_markdown_paragraph(doc, line)

    if sources:
        h = doc.add_heading("依据文献", 1)
        _apply_paragraph_font(h, 16, True)
        for x in sources:
            no = _sanitize_text(x.get("编号", ""))
            name = _sanitize_text(x.get("题名", ""))
            year = _sanitize_text(x.get("年份", ""))
            doi = _sanitize_text(x.get("DOI", ""))
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(f"[{no}] {name} ({year})" + (f"  DOI: {doi}" if doi else ""))
            _set_run_font(r, size=10.5)
            p.paragraph_format.space_after = Pt(3)

    b = BytesIO()
    doc.save(b)
    return b.getvalue()
